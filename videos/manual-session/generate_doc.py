# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# 文档路径设置
base_dir = r'D:\filework\excel-to-diagram\videos\manual-session'
doc_path = os.path.join(base_dir, '培训介绍.docx')
screenshots_dir = os.path.join(base_dir, 'screenshots')

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 标题
title = doc.add_heading('AA图 - 业务对象关系图生成工具', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 产品介绍
doc.add_heading('产品介绍', level=1)
doc.add_paragraph('AA图是一款基于Excel数据生成业务对象关系图的工具，支持Mermaid图表导出，可视化展示业务领域、对象及其关系。')

# 功能特性
doc.add_heading('功能特性', level=1)
features = [
    '[DECORATIVE] 数据导入 - 支持Excel文件导入业务对象关系数据',
    '[SEARCH] 数据校验 - 自动校验数据完整性和引用关系',
    '[SYMBOL] 领域选择 - 支持多层级业务领域选择（云领域，子域）',
    '[DECORATIVE] 多种图表类型 - 业务对象图、服务模块图、汇总图等',
    '[DESIGN] 颜色维度 - 支持按领域、类型、业务对象等维度着色',
    '[SYMBOL]️ Mermaid导出 - 支持简洁版和彩色版HTML导出',
    '[SYMBOL] 缩放拖拽 - 支持滚轮缩放和鼠标拖拽移动',
    '[SYMBOL] 视频录制 - 支持操作过程录屏',
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 操作流程
doc.add_heading('操作流程', level=1)

# 第一步
doc.add_heading('第一步：上传Excel文件', level=2)
p = doc.add_paragraph('1. 打开应用首页，点击「AA图」进入')
p = doc.add_paragraph('2. 点击上传区域，选择Excel文件')
p = doc.add_paragraph('3. 等待数据解析和校验')
p = doc.add_paragraph('说明：上传后系统自动解析数据，右侧表格显示业务对象关系明细，下方显示校验结果（错误和警告数量）。')

img_path = os.path.join(screenshots_dir, 'screenshot_001.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：上传Excel')

# 第二步
doc.add_heading('第二步：选择业务领域', level=2)
p = doc.add_paragraph('1. 在左侧树形结构中选择目标领域（如：供应链云）')
p = doc.add_paragraph('2. 展开领域节点，选择具体的子域或业务对象（如：采购供应）')
p = doc.add_paragraph('3. 选择后右侧显示该领域的业务对象和服务模块列表')
p = doc.add_paragraph('说明：支持多选，可同时选择多个业务领域进行关系分析。')

img_path = os.path.join(screenshots_dir, 'screenshot_009.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：选择领域')

# 第三步
doc.add_heading('第三步：展开并全选关系', level=2)
p = doc.add_paragraph('1. 点击「展开」按钮查看所有可选关系')
p = doc.add_paragraph('2. 点击「全选」选中所有关系')
p = doc.add_paragraph('3. 确认业务对象和服务模块列表')

img_path = os.path.join(screenshots_dir, 'screenshot_013.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：展开全选')

img_path = os.path.join(screenshots_dir, 'screenshot_015.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：全选结果')

# 第四步
doc.add_heading('第四步：选择图表类型', level=2)
p = doc.add_paragraph('1. 点击「下一步」进入图表类型选择')
p = doc.add_paragraph('2. 选择需要的图表类型：')
p = doc.add_paragraph('   • 业务对象图 - 展示业务对象之间的关系', style='List Bullet')
p = doc.add_paragraph('   • 服务模块图 - 展示服务模块及其关联', style='List Bullet')
p = doc.add_paragraph('   • 汇总图 - 高级汇总视图', style='List Bullet')
p = doc.add_paragraph('提示：双击图表类型可快速选择')

img_path = os.path.join(screenshots_dir, 'screenshot_017.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：图表类型')

# 第五步
doc.add_heading('第五步：配置图表参数', level=2)
p = doc.add_paragraph('1. 进入配置页面，设置图表参数')
p = doc.add_paragraph('2. 选择颜色维度（按领域、类型、业务对象等）')
p = doc.add_paragraph('3. 选择形状维度')
p = doc.add_paragraph('4. 点击「生成图表」预览')

img_path = os.path.join(screenshots_dir, 'screenshot_023.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：配置参数')

# 第六步
doc.add_heading('第六步：查看生成结果', level=2)
p = doc.add_paragraph('1. 系统生成Mermaid关系图')
p = doc.add_paragraph('2. 在画布区域显示图表')
p = doc.add_paragraph('3. 支持以下操作：')
p = doc.add_paragraph('   • [SYMBOL]️ 滚轮缩放 - 鼠标滚轮放大/缩小', style='List Bullet')
p = doc.add_paragraph('   • [SYMBOL] 拖拽移动 - 按住鼠标拖动画布', style='List Bullet')
p = doc.add_paragraph('   • [SYMBOL]️ 全屏查看 - 点击全屏按钮', style='List Bullet')

img_path = os.path.join(screenshots_dir, 'screenshot_025.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：图表生成')

# 第七步
doc.add_heading('第七步：导出HTML', level=2)
p = doc.add_paragraph('1. 点击导出按钮，选择导出格式：')
p = doc.add_paragraph('   • 彩色完整版 - 支持ELK布局，CDN加载，交互功能完整', style='List Bullet')
p = doc.add_paragraph('   • 简洁版 - 离线可用，内嵌Mermaid', style='List Bullet')
p = doc.add_paragraph('提示：彩色版支持更好的颜色区分和布局算法，适合复杂关系展示。')

img_path = os.path.join(screenshots_dir, 'screenshot_028.png')
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph('图：最终结果')

# 数据格式要求
doc.add_heading('数据格式要求', level=1)
p = doc.add_paragraph('Excel文件需包含以下Sheet：')

doc.add_heading('业务对象 Sheet', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '列名'
hdr_cells[1].text = '说明'
hdr_cells[2].text = '示例'
data = [
    ('编码', '业务对象唯一标识', 'BO001'),
    ('名称', '业务对象名称', '采购订单'),
    ('类型', '对象类型', '主数据/交易对象'),
    ('领域', '所属领域', '供应链云'),
]
for i, (col, desc, example) in enumerate(data):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = col
    row_cells[1].text = desc
    row_cells[2].text = example

doc.add_heading('关系 Sheet', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '列名'
hdr_cells[1].text = '说明'
hdr_cells[2].text = '示例'
data = [
    ('关系类型', '关系描述', '采购订单-库存转移'),
    ('源业务对象编码', '起始对象', 'PO001'),
    ('目标业务对象编码', '目标对象', 'INV001'),
    ('服务模块', '所属服务', '采购模块'),
]
for i, (col, desc, example) in enumerate(data):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = col
    row_cells[1].text = desc
    row_cells[2].text = example

# 使用场景
doc.add_heading('使用场景', level=1)
scenarios = [
    '业务建模 - 可视化现有业务对象及其关系',
    '系统分析 - 理解业务模块间的依赖关系',
    '文档生成 - 导出图表用于文档和汇报',
    '培训演示 - 展示业务架构全景',
]
for scenario in scenarios:
    doc.add_paragraph(scenario, style='List Bullet')

# 注意事项
doc.add_heading('注意事项', level=1)
notes = [
    '[WARNING] 数据校验 - 上传后请关注校验结果，确保数据质量',
    '[WARNING] 选择范围 - 数据量较大时注意选择范围，避免生成过多节点',
    '[WARNING] 浏览器兼容 - 推荐使用Chrome/Firefox最新版',
    '[WARNING] ELK布局 - 仅彩色完整版支持ELK布局算法',
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# 视频演示
doc.add_heading('视频演示', level=1)
doc.add_paragraph('完整操作演示视频：page@bb5ddf60a680a91ad939b2961bce4861.webm')

# 保存文档
doc.save(doc_path)
print(f'Word文档已生成: {doc_path}')
